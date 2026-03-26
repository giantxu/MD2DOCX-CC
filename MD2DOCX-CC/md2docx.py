#!/usr/bin/env python3
"""
md2docx.py
将 Markdown 法律文书转换为符合模板排版规范的 Word 文档。

用法:
    python3 md2docx.py <input.md> <output.docx> [template.docx]

编号定义由脚本自动创建，不依赖模板中已有的 numId，任何 docx 模板均可使用。
模板仅提供：页面尺寸/页边距、页眉页脚、Normal / List Paragraph 等段落样式的
默认字体与字号。
"""

import re
import sys
import random
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 间距常量（pt）────────────────────────────────────────────────────────────
TITLE_PT        = 14
TITLE_SPC_BEF   = 18
TITLE_SPC_AFT   = 18
BODY_SPC_AFT    = 6
SEC_SPC_BEF     = 5
SEC_SPC_AFT     = 5
SUBSEC_SPC_BEF  = 15
SUBSEC_SPC_AFT  = 9


# ════════════════════════════════════════════════════════════════════════════
# 编号定义：自动创建，不依赖模板已有 numId
# ════════════════════════════════════════════════════════════════════════════

def _random_nsid() -> str:
    """生成 8 位十六进制 nsid（Word 要求每个 abstractNum 有唯一 nsid）。"""
    return f'{random.randint(0, 0xFFFFFFFF):08X}'


def _make_lvl(ilvl: int, num_fmt: str, lvl_text: str,
              left: int, hanging: int, bold: bool = False,
              font_hint: str | None = None) -> OxmlElement:
    """构建一个 <w:lvl> 元素。"""
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), str(ilvl))

    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    lvl.append(start)

    fmt = OxmlElement('w:numFmt')
    fmt.set(qn('w:val'), num_fmt)
    lvl.append(fmt)

    txt = OxmlElement('w:lvlText')
    txt.set(qn('w:val'), lvl_text)
    lvl.append(txt)

    jc = OxmlElement('w:lvlJc')
    jc.set(qn('w:val'), 'left')
    lvl.append(jc)

    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left))
    ind.set(qn('w:hanging'), str(hanging))
    pPr.append(ind)
    lvl.append(pPr)

    if bold or font_hint:
        rPr = OxmlElement('w:rPr')
        if bold:
            rPr.append(OxmlElement('w:b'))
        if font_hint:
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:hint'), font_hint)
            rPr.append(rf)
        lvl.append(rPr)

    return lvl


def _make_abstract_num(abs_id: int, multi_type: str,
                       levels: list[OxmlElement]) -> OxmlElement:
    """构建一个 <w:abstractNum> 元素。"""
    an = OxmlElement('w:abstractNum')
    an.set(qn('w:abstractNumId'), str(abs_id))

    nsid = OxmlElement('w:nsid')
    nsid.set(qn('w:val'), _random_nsid())
    an.append(nsid)

    mlt = OxmlElement('w:multiLevelType')
    mlt.set(qn('w:val'), multi_type)
    an.append(mlt)

    tmpl = OxmlElement('w:tmpl')
    tmpl.set(qn('w:val'), _random_nsid())
    an.append(tmpl)

    for lvl in levels:
        an.append(lvl)
    return an


def _make_num(num_id: int, abs_id: int, start_override: int | None = None) -> OxmlElement:
    """构建一个 <w:num> 元素，可选 startOverride。"""
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))

    ref = OxmlElement('w:abstractNumId')
    ref.set(qn('w:val'), str(abs_id))
    num.append(ref)

    if start_override is not None:
        ov = OxmlElement('w:lvlOverride')
        ov.set(qn('w:ilvl'), '0')
        so = OxmlElement('w:startOverride')
        so.set(qn('w:val'), str(start_override))
        ov.append(so)
        num.append(ov)

    return num


class NumConfig:
    """保存由 _setup_numbering 创建的各类编号 numId / abstractNumId。"""
    __slots__ = ('section_nid', 'subsect_nid', 'bullet_nid',
                 'ordered_abs_id', '_numbering', '_max_num_id')

    def __init__(self):
        self.section_nid: int = 0
        self.subsect_nid: int = 0
        self.bullet_nid: int = 0
        self.ordered_abs_id: int = 0
        self._numbering = None
        self._max_num_id: int = 0

    def create_ordered_num(self) -> int:
        """为每组独立有序列表克隆一个新 numId（从 1 开始）。"""
        self._max_num_id += 1
        num = _make_num(self._max_num_id, self.ordered_abs_id, start_override=1)
        self._numbering.append(num)
        return self._max_num_id


def _get_or_create_numbering_part(doc):
    """获取或创建 numbering part，返回其 XML 根元素。"""
    try:
        return doc.part.numbering_part._element
    except Exception:
        # 模板中不存在 numbering part，用 python-docx 内部方法创建
        # 先添加一个临时列表段落以触发 numbering part 的创建
        tmp = doc.add_paragraph()
        tmp.style = doc.styles['List Paragraph']
        # 触发创建后删除临时段落
        tmp._element.getparent().remove(tmp._element)
        return doc.part.numbering_part._element


def _setup_numbering(doc) -> NumConfig:
    """
    在文档中创建所有需要的编号定义（abstractNum + num），
    完全不依赖模板已有的 numId/abstractNumId。
    """
    numbering = _get_or_create_numbering_part(doc)
    cfg = NumConfig()
    cfg._numbering = numbering

    # ── 计算当前最大 ID，避免冲突 ────────────────────────────────────────
    existing_abs = numbering.findall(qn('w:abstractNum'))
    existing_nums = numbering.findall(qn('w:num'))
    max_abs = max((int(a.get(qn('w:abstractNumId'))) for a in existing_abs), default=0)
    max_num = max((int(n.get(qn('w:numId'))) for n in existing_nums), default=0)

    next_abs = max_abs + 1
    next_num = max_num + 1

    # ── 找到插入点：abstractNum 必须在所有 num 之前 ────────────────────────
    first_num = numbering.find(qn('w:num'))

    def _insert_abstract(an):
        if first_num is not None:
            first_num.addprevious(an)
        else:
            numbering.append(an)

    # ── 1. 中文数字编号（一、二、三…）→ ## 大节标题 ─────────────────────────
    abs_section = _make_abstract_num(next_abs, 'hybridMultilevel', [
        _make_lvl(0, 'chineseCountingThousand', '%1\u3001',
                  left=420, hanging=420, bold=True),
    ])
    _insert_abstract(abs_section)
    num_section = _make_num(next_num, next_abs)
    numbering.append(num_section)
    cfg.section_nid = next_num
    next_abs += 1; next_num += 1

    # ── 2. 多级 decimal（1. / 1.1 / 1.1.1）→ ### 小节标题 + 子项 ──────────
    abs_subsect = _make_abstract_num(next_abs, 'multilevel', [
        _make_lvl(0, 'decimal', '%1.',
                  left=425, hanging=425, font_hint='eastAsia'),
        _make_lvl(1, 'decimal', '%1.%2',
                  left=992, hanging=567, font_hint='eastAsia'),
        _make_lvl(2, 'decimal', '%1.%2.%3',
                  left=1418, hanging=567, font_hint='eastAsia'),
    ])
    _insert_abstract(abs_subsect)
    num_subsect = _make_num(next_num, next_abs)
    numbering.append(num_subsect)
    cfg.subsect_nid = next_num
    next_abs += 1; next_num += 1

    # ── 3. 项目符号列表（●  ○  ■）→ * 无序列表 ─────────────────────────────
    abs_bullet = _make_abstract_num(next_abs, 'hybridMultilevel', [
        _make_lvl(0, 'bullet', '\u25CF',   # ●
                  left=720, hanging=360),
        _make_lvl(1, 'bullet', '\u25CB',    # ○
                  left=1440, hanging=360),
        _make_lvl(2, 'bullet', '\u25A0',    # ■
                  left=2160, hanging=360),
    ])
    _insert_abstract(abs_bullet)
    num_bullet = _make_num(next_num, next_abs)
    numbering.append(num_bullet)
    cfg.bullet_nid = next_num
    next_abs += 1; next_num += 1

    # ── 4. 带加粗的 decimal（1.  / 1.1）→ 有序列表正文项 ────────────────────
    abs_ordered = _make_abstract_num(next_abs, 'multilevel', [
        _make_lvl(0, 'decimal', '%1.',
                  left=425, hanging=425, bold=True, font_hint='eastAsia'),
        _make_lvl(1, 'decimal', '%1.%2',
                  left=992, hanging=567, font_hint='eastAsia'),
        _make_lvl(2, 'decimal', '%1.%2.%3',
                  left=1418, hanging=567, font_hint='eastAsia'),
    ])
    _insert_abstract(abs_ordered)
    cfg.ordered_abs_id = next_abs
    next_abs += 1

    cfg._max_num_id = next_num - 1
    return cfg


# ════════════════════════════════════════════════════════════════════════════
# 段落编号设置
# ════════════════════════════════════════════════════════════════════════════

def _set_numbering(para, num_id: int, ilvl: int):
    """为段落设置 Word 自动编号。"""
    pPr = para._element.get_or_add_pPr()
    for old in pPr.findall(qn('w:numPr')):
        pPr.remove(old)
    numPr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_el)
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))
    numPr.append(numId_el)
    pPr.append(numPr)


# ════════════════════════════════════════════════════════════════════════════
# 行内格式解析
# ════════════════════════════════════════════════════════════════════════════

# 匹配 **bold** 或 *italic*（不跨行）
_INLINE_PAT = re.compile(r'\*\*(.+?)\*\*|\*([^*\n]+)\*')


def parse_inline(text: str) -> list[tuple[str, bool, bool]]:
    """
    解析行内 Markdown 格式，返回 [(text, bold, italic), ...] 列表。
    - **text** → bold=True
    - *text*   → italic=True
    - <span>…</span> 等 HTML 标签保留为原文
    """
    result = []
    last = 0
    for m in _INLINE_PAT.finditer(text):
        if m.start() > last:
            result.append((text[last:m.start()], False, False))
        if m.group(0).startswith('**'):
            result.append((m.group(1), True, False))
        else:
            result.append((m.group(2), False, True))
        last = m.end()
    if last < len(text):
        result.append((text[last:], False, False))
    return result


def _add_runs(para, text: str):
    """将带行内格式的文本添加到段落。"""
    for seg_text, bold, italic in parse_inline(text):
        if not seg_text:
            continue
        run = para.add_run(seg_text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True


# ════════════════════════════════════════════════════════════════════════════
# 标题文本提取
# ════════════════════════════════════════════════════════════════════════════

# 匹配 ATX 标题：1–6 个 # 后跟至少一个空格
_HEADING_RE = re.compile(r'^(#{1,6})\s+(.+?)(?:\s+#+\s*)?$')

# 中文数字前缀（含复合数字如"十一"、"二十三"等）
_CN_NUM_PREFIX = re.compile(
    r'^[零一二三四五六七八九十百千]+[、．.]\s*'
)
# 阿拉伯数字前缀
_DIGIT_PREFIX = re.compile(r'^\d+[\.．、)）]\s*')


def _strip_heading_prefix(text: str) -> str:
    """去掉标题中手动写的序号前缀（中文数字或阿拉伯数字）。"""
    text = _CN_NUM_PREFIX.sub('', text)
    text = _DIGIT_PREFIX.sub('', text)
    return text


# ════════════════════════════════════════════════════════════════════════════
# 段落构建器
# ════════════════════════════════════════════════════════════════════════════

def _spc(para, before_pt=None, after_pt=None):
    fmt = para.paragraph_format
    if before_pt is not None:
        fmt.space_before = Pt(before_pt)
    if after_pt is not None:
        fmt.space_after = Pt(after_pt)


def add_title(doc, text: str, _nc=None):
    """# 一级标题 → 居中加粗 14pt。"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spc(para, TITLE_SPC_BEF, TITLE_SPC_AFT)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(TITLE_PT)


def add_section_heading(doc, text: str, nc: NumConfig):
    """## 大节标题 → List Paragraph + 中文数字自动编号（一、二、三…）。"""
    text = _strip_heading_prefix(text)
    para = doc.add_paragraph()
    para.style = doc.styles['List Paragraph']
    _spc(para, SEC_SPC_BEF, SEC_SPC_AFT)
    _set_numbering(para, nc.section_nid, 0)
    run = para.add_run(text)
    run.bold = True


def add_subsection_heading(doc, text: str, nc: NumConfig):
    """### 小节标题 → List Paragraph + decimal 自动编号 ilvl=0。"""
    text = _strip_heading_prefix(text)
    para = doc.add_paragraph()
    para.style = doc.styles['List Paragraph']
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spc(para, SUBSEC_SPC_BEF, SUBSEC_SPC_AFT)
    _set_numbering(para, nc.subsect_nid, 0)
    run = para.add_run(text)
    run.bold = True


def add_h4_heading(doc, text: str, _nc=None):
    """#### 四级标题 → List Paragraph 加粗，不编号。"""
    text = _strip_heading_prefix(text)
    para = doc.add_paragraph()
    para.style = doc.styles['List Paragraph']
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spc(para, after_pt=BODY_SPC_AFT)
    run = para.add_run(text)
    run.bold = True


def add_h5_heading(doc, text: str, _nc=None):
    """##### 五级及更深 → 加粗斜体正文段落。"""
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spc(para, after_pt=BODY_SPC_AFT)
    run = para.add_run(text)
    run.bold = True
    run.italic = True


def add_body_para(doc, text: str):
    """正文段落 → Normal 样式，两端对齐，段后 6pt。"""
    if not text.strip():
        return
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    _spc(para, after_pt=BODY_SPC_AFT)
    _add_runs(para, text)


def add_bullet_item(doc, text: str, nc: NumConfig):
    """无序列表项（* 纯文本项）→ List Paragraph + bullet ilvl=0。"""
    para = doc.add_paragraph()
    para.style = doc.styles['List Paragraph']
    _spc(para, after_pt=BODY_SPC_AFT)
    _set_numbering(para, nc.bullet_nid, 0)
    _add_runs(para, text)


def add_ordered_sub_item(doc, text: str, nc: NumConfig):
    """修改建议类有序子项（* **label:** 文本）→ subsect 编号 ilvl=1。"""
    para = doc.add_paragraph()
    para.style = doc.styles['List Paragraph']
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spc(para, after_pt=BODY_SPC_AFT)
    _set_numbering(para, nc.subsect_nid, 1)
    _add_runs(para, text)


def add_numbered_item(doc, text: str, num_id: int, ilvl: int = 0):
    """有序列表项（1. 2. 3. 或其子项）→ List Paragraph + 动态 numId。"""
    para = doc.add_paragraph()
    para.style = doc.styles['List Paragraph']
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spc(para, after_pt=BODY_SPC_AFT)
    _set_numbering(para, num_id, ilvl)
    _add_runs(para, text)


# ════════════════════════════════════════════════════════════════════════════
# 表格构建器
# ════════════════════════════════════════════════════════════════════════════

def _parse_table_row(line: str) -> list[str]:
    """将 | col | col | 行拆分为单元格列表。"""
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


def _apply_table_borders(table):
    """为表格添加全边框（单线 0.5pt）。"""
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._element.insert(0, tblPr)

    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)

    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)

    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)


def add_table(doc, raw_rows: list[list[str]]):
    """添加表格。表头行加粗居中，<br> 拆分为单元格内多段落。"""
    if not raw_rows:
        return

    data_rows = []
    for row in raw_rows:
        if all(re.match(r'^[-:\s]*$', c) for c in row):
            continue
        data_rows.append(row)

    if not data_rows:
        return

    n_cols = len(data_rows[0])
    table = doc.add_table(rows=len(data_rows), cols=n_cols)
    table.style = 'Normal Table'
    _apply_table_borders(table)

    for i, row_cells in enumerate(data_rows):
        row = table.rows[i]
        is_header = (i == 0)
        for j, cell_text in enumerate(row_cells[:n_cols]):
            cell = row.cells[j]
            parts = re.split(r'<br\s*/?>', cell_text)
            first_para = cell.paragraphs[0]
            if is_header:
                first_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_runs(first_para, parts[0])
            if is_header:
                for run in first_para.runs:
                    run.bold = True
            for part in parts[1:]:
                new_para = cell.add_paragraph()
                _add_runs(new_para, part)


# ════════════════════════════════════════════════════════════════════════════
# 模板加载
# ════════════════════════════════════════════════════════════════════════════

def load_clean_template(template_path: str) -> Document:
    """打开模板 docx，清空正文内容，保留页面设置 / 页眉页脚 / 样式。"""
    doc = Document(template_path)
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl', 'sdt'):
            body.remove(child)
    return doc


# ════════════════════════════════════════════════════════════════════════════
# 行类型检测
# ════════════════════════════════════════════════════════════════════════════

def _is_numbered_item(line: str) -> bool:
    """行首为 `数字.` 或 `数字．` 后跟空格 → 有序列表项。"""
    return bool(re.match(r'^\d+[\.．]\s', line.strip()))


def _is_sub_bullet(line: str) -> bool:
    """缩进子项：以 2+ 空格（或 tab）开头 + * 或 - 符号。"""
    return bool(re.match(r'^(?:  +|\t+)[*-]\s', line))


def _ends_ordered_list(line: str) -> bool:
    """判断该行是否会终止当前有序列表（空行、列表子项不终止）。"""
    stripped = line.strip()
    if not stripped:
        return False
    if _is_numbered_item(line):
        return False
    if _is_sub_bullet(line):
        return False
    return True


def _next_non_blank(lines: list[str], start: int) -> str:
    """返回 start 之后第一个非空行。"""
    for k in range(start + 1, len(lines)):
        if lines[k].strip():
            return lines[k]
    return ''


# ════════════════════════════════════════════════════════════════════════════
# 主解析器
# ════════════════════════════════════════════════════════════════════════════

def convert(md_path: str, out_path: str, template_path: str = 'Template.docx'):
    doc = load_clean_template(template_path)
    nc  = _setup_numbering(doc)

    text = Path(md_path).read_text(encoding='utf-8')
    lines = text.splitlines()

    i = 0
    ordered_nid = None
    in_ordered  = False

    while i < len(lines):
        raw  = lines[i]
        line = raw.strip()

        # ── 空行 ────────────────────────────────────────────────────────────
        if not line:
            next_line = _next_non_blank(lines, i)
            if in_ordered and _ends_ordered_list(next_line):
                in_ordered = False
                ordered_nid = None
            i += 1
            continue

        # ── 表格（| 开头）──────────────────────────────────────────────────
        if line.startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_rows.append(_parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            in_ordered = False; ordered_nid = None
            continue

        # ── 分隔线（--- / *** / ___）────────────────────────────────────────
        if re.match(r'^[-*_]{3,}\s*$', line):
            i += 1
            continue

        # ── ATX 标题 ────────────────────────────────────────────────────────
        hm = _HEADING_RE.match(line)
        if hm:
            level = len(hm.group(1))
            h_text = hm.group(2).strip()
            in_ordered = False; ordered_nid = None
            if level == 1:
                add_title(doc, h_text)
            elif level == 2:
                add_section_heading(doc, h_text, nc)
            elif level == 3:
                add_subsection_heading(doc, h_text, nc)
            elif level == 4:
                add_h4_heading(doc, h_text)
            else:
                add_h5_heading(doc, h_text)
            i += 1
            continue

        # ── 有序列表项（1. 2. 3.）───────────────────────────────────────────
        if _is_numbered_item(raw):
            if not in_ordered:
                ordered_nid = nc.create_ordered_num()
                in_ordered  = True
            item_text = re.sub(r'^\d+[\.．]\s+', '', line)
            add_numbered_item(doc, item_text, ordered_nid, ilvl=0)
            i += 1
            continue

        # ── 缩进子项（有序列表的 * / - 子弹）──────────────────────────────
        if _is_sub_bullet(raw):
            item_text = re.sub(r'^\s+[*-]\s+', '', raw)
            if in_ordered and ordered_nid:
                add_numbered_item(doc, item_text, ordered_nid, ilvl=1)
            else:
                add_bullet_item(doc, item_text, nc)
            i += 1
            continue

        # ── 无序列表项（* 或 - 开头）────────────────────────────────────────
        if re.match(r'^[*-]\s+', line):
            item_text = re.sub(r'^[*-]\s+', '', line)
            # 以 **label:** 开头 → 修改建议类有序子项
            if re.match(r'^\*\*[^*]+[：:]\*\*', item_text):
                add_ordered_sub_item(doc, item_text, nc)
            else:
                add_bullet_item(doc, item_text, nc)
            i += 1
            continue

        # ── 正文段落 ─────────────────────────────────────────────────────────
        if in_ordered and _ends_ordered_list(raw):
            in_ordered = False
            ordered_nid = None
        add_body_para(doc, line)
        i += 1

    doc.save(out_path)
    print(f'已生成：{out_path}')


# ════════════════════════════════════════════════════════════════════════════
# 入口
# ════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('用法: python3 md2docx.py <input.md> <output.docx> [template.docx]')
        sys.exit(1)

    md_in    = sys.argv[1]
    docx_out = sys.argv[2]
    tmpl     = sys.argv[3] if len(sys.argv) > 3 else str(
        Path(__file__).parent / 'Template.docx'
    )
    convert(md_in, docx_out, tmpl)
